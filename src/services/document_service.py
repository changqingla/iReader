"""Document service business logic."""
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from fastapi import HTTPException, status, UploadFile
from repositories.kb_repository import KnowledgeBaseRepository
from repositories.document_repository import DocumentRepository
from repositories.user_repository import UserRepository
from repositories.organization_member_repository import OrganizationMemberRepository
from utils.minio_client import upload_file, delete_file
from utils.external_services import MineruService, DocumentProcessService
from utils.es_utils import get_user_es_index
from models.document import Document
from models.knowledge_base import KnowledgeBase
from config.settings import settings
from typing import List, Tuple, Optional
import os
import logging
import asyncio
import uuid

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document operations."""
    
    # Supported file extensions (知识库只支持这5种格式)
    PDF_EXTENSIONS = {'.pdf'}
    TEXT_EXTENSIONS = {'.md', '.markdown', '.txt'}
    WORD_EXTENSIONS = {'.doc', '.docx'}
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.kb_repo = KnowledgeBaseRepository(db)
        self.doc_repo = DocumentRepository(db)
        self.user_repo = UserRepository(db)
        self.org_member_repo = OrganizationMemberRepository(db)
    
    async def _verify_kb_write_access(self, kb_id: str, user_id: str) -> KnowledgeBase:
        """
        Verify user has WRITE access to knowledge base.
        Only owner and admin users have write permissions.
        
        Returns:
            Knowledge base object if user has write access
        
        Raises:
            HTTPException: If knowledge base not found or user has no write permission
        """
        # Check if user is admin
        user = await self.user_repo.get_by_id(uuid.UUID(user_id))
        is_admin = user and user.is_admin
        
        # Try to get as owner
        kb = await self.kb_repo.get_by_id(kb_id, user_id)
        
        # If not owner but is admin, get the KB
        if not kb and is_admin:
            kb = await self.kb_repo.get_by_id_any(kb_id)
        
        if not kb:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail={"error": {"code": "FORBIDDEN", "message": "Only the knowledge base owner or admin can perform this action"}}
            )
        
        return kb
    
    async def _verify_kb_access(self, kb_id: str, user_id: str) -> KnowledgeBase:
        """
        Verify user has access to knowledge base.
        - Admin users: can access any knowledge base
        - Owners: can access their own knowledge bases
        - Organization members: can access organization-shared knowledge bases
        - Everyone: can access public knowledge bases
        
        Returns:
            Knowledge base object if accessible
        
        Raises:
            HTTPException: If knowledge base not found or not accessible
        """
        # Check if user is admin
        user = await self.user_repo.get_by_id(uuid.UUID(user_id))
        is_admin = user and user.is_admin
        
        # Try to get as owner first
        kb = await self.kb_repo.get_by_id(kb_id, user_id)
        
        # If not owner and admin, get any KB
        if not kb and is_admin:
            kb = await self.kb_repo.get_by_id_any(kb_id)
        
        # If not owner and not admin, check organization-shared or public KB
        if not kb:
            # Get user's organizations
            user_org_ids = await self.org_member_repo.get_user_org_ids(uuid.UUID(user_id))
            
            # Try to get the KB without access check
            kb = await self.kb_repo.get_by_id_any(kb_id)
            
            if kb:
                # Check if user has access
                has_access = False
                
                # 1. Public KB - everyone can access
                if kb.visibility == 'public':
                    has_access = True
                
                # 2. Organization-shared KB - check if user is in any shared organization
                elif kb.visibility == 'organization' and user_org_ids:
                    # Check if any of user's org IDs is in the shared_to_orgs
                    shared_org_ids = set(kb.shared_to_orgs or [])
                    user_orgs_set = set(user_org_ids)
                    if shared_org_ids.intersection(user_orgs_set):
                        has_access = True
                
                if not has_access:
                    kb = None
            
            if not kb:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail={"error": {"code": "NOT_FOUND", "message": "Knowledge base not found or not accessible"}}
                )
        
        return kb
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension in lowercase."""
        return os.path.splitext(filename)[1].lower()
    
    def _needs_mineru_conversion(self, filename: str) -> bool:
        """Check if file needs Mineru conversion (PDF only)."""
        ext = self._get_file_extension(filename)
        return ext in self.PDF_EXTENSIONS
    
    def _is_word_document(self, filename: str) -> bool:
        """Check if file is a Word document (.doc or .docx)."""
        ext = self._get_file_extension(filename)
        return ext in self.WORD_EXTENSIONS
    
    def _is_text_file(self, filename: str) -> bool:
        """Check if file is a plain text file (.txt, .md, .markdown)."""
        ext = self._get_file_extension(filename)
        return ext in self.TEXT_EXTENSIONS
    
    def _extract_docx_content(self, file_data: bytes) -> str:
        """
        Extract text content from .docx file.
        First tries python-docx, falls back to Tika if that fails.
        
        Args:
            file_data: Binary content of the docx file
            
        Returns:
            Extracted text content as string
        """
        from io import BytesIO
        
        # First, try python-docx
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(BytesIO(file_data))
            paragraphs = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            content = '\n\n'.join(paragraphs)
            if content:
                return content
            
            # If python-docx returned empty, try Tika
            logger.warning("python-docx returned empty content, trying Tika...")
            
        except Exception as e:
            # python-docx failed, try Tika as fallback
            logger.warning(f"python-docx failed: {e}, trying Tika as fallback...")
        
        # Fallback to Tika
        try:
            from tika import parser
            result = parser.from_buffer(BytesIO(file_data))
            content = result.get('content', '')
            
            if content:
                lines = [line.strip() for line in content.split('\n') if line.strip()]
                return '\n\n'.join(lines)
        except Exception as e:
            logger.error(f"Tika also failed for DOCX: {e}")
        
        return ''
    
    def _extract_doc_content(self, file_data: bytes) -> str:
        """
        Extract text content from .doc file using Apache Tika.
        
        Args:
            file_data: Binary content of the doc file
            
        Returns:
            Extracted text content as string
            
        Raises:
            Exception: If Tika fails to extract content
        """
        from tika import parser
        from io import BytesIO
        
        result = parser.from_buffer(BytesIO(file_data))
        content = result.get('content', '')
        
        if content:
            # Clean up the content (remove excessive whitespace)
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            return '\n\n'.join(lines)
        
        # Return empty string, caller should handle this case
        return ''
    
    async def upload_document(
        self,
        kb_id: str,
        user_id: str,
        file: UploadFile,
        background_tasks
    ) -> dict:
        """
        Upload document and trigger processing.
        Only owner and admin users can upload documents.
        
        Supported formats: pdf, txt, md, doc, docx
        
        Complete flow:
        1. Upload to MinIO
        2. Create document record
        3. If PDF: convert with Mineru
        4. Process document (chunk + embed + store to ES)
        5. Background task to poll status
        """
        # Verify KB write access (owner or admin only)
        kb = await self._verify_kb_write_access(kb_id, user_id)
        
        # Validate file format
        ext = self._get_file_extension(file.filename)
        supported_extensions = self.PDF_EXTENSIONS | self.TEXT_EXTENSIONS | self.WORD_EXTENSIONS
        if ext not in supported_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": {"code": "UNSUPPORTED_FORMAT", "message": f"Unsupported file format: {ext}. Supported formats: pdf, txt, md, doc, docx"}}
            )
        
        # Read file
        file_data = await file.read()
        file_size = len(file_data)
        
        # Upload to MinIO
        object_name = f"kb/{user_id}/{kb_id}/{file.filename}"
        try:
            file_path = await upload_file(object_name, file_data, file.content_type)
        except Exception as e:
            logger.error(f"MinIO upload failed: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": {"code": "INTERNAL_ERROR", "message": f"File upload failed: {e}"}}
            )
        
        # Create document record
        document = await self.doc_repo.create(
            kb_id=kb_id,
            name=file.filename,
            size=file_size,
            source="upload",
            file_path=file_path
        )
        
        # Get user's ES index name (user-level, not KB-level)
        user_es_index = get_user_es_index(user_id)
        logger.info(f"Using ES index: {user_es_index} for user {user_id}")
        
        # Start background processing using FastAPI BackgroundTasks
        logger.info(f"Starting background processing for document {document.id} ({file.filename})")
        background_tasks.add_task(
            self._process_document_pipeline,
            str(document.id),
            user_es_index,
            file_data,
            file.filename
        )
        
        return {
            "id": str(document.id),
            "name": document.name,
            "status": document.status
        }
    
    async def _process_document_pipeline(
        self,
        doc_id: str,
        es_index_name: str,
        file_data: bytes,
        filename: str
    ):
        """
        Background task to process document through complete pipeline.
        
        Pipeline:
        1. Convert with Mineru (if PDF)
        2. Parse document (chunk + embed + store)
        3. Update status
        """
        # Import here to avoid circular dependency
        from config.database import AsyncSessionLocal
        from repositories.document_repository import DocumentRepository
        from repositories.kb_repository import KnowledgeBaseRepository
        
        # Create new DB session for background task
        async with AsyncSessionLocal() as db:
            doc_repo = DocumentRepository(db)
            kb_repo = KnowledgeBaseRepository(db)
            
            # Get document
            result = await db.execute(
                select(Document).where(Document.id == doc_id)
            )
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error(f"Document {doc_id} not found in background task")
                return
            
            logger.info(f"[Doc {doc_id}] Background task started for {filename}")
        
            try:
                markdown_content = None
                ext = self._get_file_extension(filename)
                
                # Step 1: Extract content based on file type
                if self._needs_mineru_conversion(filename):
                    # PDF files: use MinerU for conversion
                    logger.info(f"[Doc {doc_id}] PDF detected, calling Mineru for conversion")
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)
                
                    try:
                        logger.info(f"[Doc {doc_id}] Calling MinerU official API...")
                        mineru_result = await MineruService.convert_document(file_data, filename)
                        batch_id = mineru_result["batch_id"]
                        logger.info(f"[Doc {doc_id}] MinerU task created, batch_id: {batch_id}")

                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_PROCESSING,
                            mineru_task_id=batch_id
                        )

                        # Poll MinerU status
                        logger.info(f"[Doc {doc_id}] Polling MinerU task status...")
                        markdown_content = await self._poll_mineru_task(batch_id, doc_id)
                        logger.info(f"[Doc {doc_id}] MinerU conversion completed, got {len(markdown_content)} chars")
                        
                    except Exception as e:
                        logger.error(f"[Doc {doc_id}] Mineru conversion failed: {e}")
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=f"Conversion failed: {e}"
                        )
                        return
                
                elif ext == '.docx':
                    # DOCX files: use python-docx to extract content
                    logger.info(f"[Doc {doc_id}] DOCX detected, extracting content with python-docx")
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)
                    try:
                        markdown_content = self._extract_docx_content(file_data)
                        if not markdown_content:
                            raise Exception("DOCX file appears to be empty or contains no extractable text")
                        logger.info(f"[Doc {doc_id}] DOCX content extracted, got {len(markdown_content)} chars")
                    except Exception as e:
                        logger.error(f"[Doc {doc_id}] DOCX extraction failed: {e}")
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=f"DOCX extraction failed: {e}"
                        )
                        return
                
                elif ext == '.doc':
                    # DOC files: use Apache Tika to extract content
                    logger.info(f"[Doc {doc_id}] DOC detected, extracting content with Tika")
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)
                    try:
                        markdown_content = self._extract_doc_content(file_data)
                        if not markdown_content:
                            raise Exception("Tika returned empty content")
                        logger.info(f"[Doc {doc_id}] DOC content extracted, got {len(markdown_content)} chars")
                    except Exception as e:
                        logger.error(f"[Doc {doc_id}] DOC extraction failed: {e}")
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=f"DOC extraction failed: {e}"
                        )
                        return
                
                else:
                    # Text files (.txt, .md, .markdown): read directly
                    logger.info(f"[Doc {doc_id}] Text file detected, reading directly")
                    try:
                        markdown_content = file_data.decode('utf-8')
                    except UnicodeDecodeError:
                        # Try other encodings
                        try:
                            markdown_content = file_data.decode('gbk')
                        except UnicodeDecodeError:
                            markdown_content = file_data.decode('latin-1')
                    logger.info(f"[Doc {doc_id}] Text content read, got {len(markdown_content)} chars")
                
                # Save markdown content to MinIO (for agent use)
                logger.info(f"[Doc {doc_id}] Saving content to MinIO...")
                kb_id = str(doc.kb_id)
                user_id = doc.file_path.split('/')[1] if doc.file_path else "unknown"  # Extract from file_path
                md_object_name = f"kb/{user_id}/{kb_id}/markdown/{doc_id}.md"
                
                try:
                    markdown_bytes = markdown_content.encode('utf-8')
                    markdown_path = await upload_file(md_object_name, markdown_bytes, "text/markdown")
                    await doc_repo.update_markdown_path(doc, markdown_path)
                    logger.info(f"[Doc {doc_id}] Content saved to MinIO: {markdown_path}")
                except Exception as e:
                    logger.warning(f"[Doc {doc_id}] Failed to save content to MinIO: {e}")
                    # Continue processing even if MinIO save fails
                
                # Step 2: Parse document (chunk + embed + store to ES)
                await doc_repo.update_status(doc, Document.STATUS_CHUNKING)
                
                # Always send extracted text content to rag-agent for chunking
                # This ensures consistent processing and avoids issues with problematic files
                temp_file_path = f"/tmp/{doc_id}.md"
                with open(temp_file_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                parse_filename = os.path.splitext(filename)[0] + '.md'
                logger.info(f"[Doc {doc_id}] Saved content to {temp_file_path}, will use filename: {parse_filename}")
                
                try:
                    logger.info(f"[Doc {doc_id}] Calling document processing service...")
                    parse_result = await DocumentProcessService.parse_document(
                        temp_file_path,
                        str(doc_id),
                        es_index_name,
                        parse_filename
                    )
                    task_id = parse_result["task_id"]
                    logger.info(f"[Doc {doc_id}] Parse task created: {task_id}")
                    
                    await doc_repo.update_status(
                        doc,
                        Document.STATUS_EMBEDDING,
                        parse_task_id=task_id
                    )
                    
                    # Poll parsing status
                    logger.info(f"[Doc {doc_id}] Polling parse task status...")
                    success = await self._poll_parse_task(doc, task_id, doc_repo)
                    
                    if success:
                        logger.info(f"[Doc {doc_id}] Document processing completed successfully!")
                        # Increment KB contents count only on success
                        await kb_repo.increment_contents_count(doc.kb_id)
                    else:
                        logger.warning(f"[Doc {doc_id}] Document processing failed, not incrementing count")
                    
                finally:
                    # Clean up temp file
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                
            except Exception as e:
                logger.error(f"Document processing pipeline failed for {doc_id}: {e}")
                await doc_repo.update_status(
                    doc,
                    Document.STATUS_FAILED,
                    error_message=str(e)
                )
    
    async def _poll_mineru_task(self, batch_id: str, doc_id: str = None, max_attempts: int = 180) -> str:
        """
        Poll MinerU task until completion.

        Args:
            batch_id: MinerU batch ID
            doc_id: Document ID for logging
            max_attempts: Maximum polling attempts (default: 180 = 15 minutes with 5s interval)

        Returns:
            Markdown content string
        """
        log_prefix = f"[Doc {doc_id}]" if doc_id else "[MinerU]"
        last_progress = None

        for attempt in range(max_attempts):
            await asyncio.sleep(5)  # Wait 5 seconds

            try:
                task_status = await MineruService.get_task_status(batch_id)
                status = task_status["status"]

                # 记录进度
                if status == "running":
                    progress = task_status.get("progress", {})
                    extracted = progress.get("extracted_pages", 0)
                    total = progress.get("total_pages", 0)
                    if (extracted, total) != last_progress:
                        logger.info(f"{log_prefix} MinerU processing: {extracted}/{total} pages")
                        last_progress = (extracted, total)
                elif status == "pending":
                    if attempt % 6 == 0:  # 每 30 秒记录一次
                        logger.info(f"{log_prefix} MinerU task pending, waiting...")

                if status == "completed":
                    logger.info(f"{log_prefix} MinerU task completed, downloading content...")
                    return await MineruService.get_content(batch_id)
                elif status == "failed":
                    error_msg = task_status.get('message', 'Unknown error')
                    raise Exception(f"MinerU task failed: {error_msg}")

            except Exception as e:
                if "failed" in str(e).lower():
                    raise  # 重新抛出失败异常
                logger.warning(f"{log_prefix} Error polling MinerU task {batch_id}: {e}")

        raise Exception("MinerU task timeout after 15 minutes")
    
    async def _poll_parse_task(self, doc: Document, task_id: str, doc_repo, max_attempts: int = 120) -> bool:
        """
        Poll document parsing task until completion (timeout: 10 minutes).
        
        Returns:
            True if processing succeeded, False if failed
        """
        for _ in range(max_attempts):
            await asyncio.sleep(5)  # Wait 5 seconds
            
            try:
                task_status = await DocumentProcessService.get_task_status(task_id)
                status = task_status["status"]
                
                if status == "completed":
                    # Get chunk count from task data
                    chunk_count = task_status.get("data", {}).get("total_chunks", 0)
                    await doc_repo.update_status(
                        doc,
                        Document.STATUS_READY,
                        chunk_count=chunk_count
                    )
                    logger.info(f"Document {doc.id} processing completed with {chunk_count} chunks")
                    return True
                
                elif status == "failed":
                    error_msg = task_status.get("message", "Processing failed")
                    logger.error(f"[Doc {doc.id}] Parse task {task_id} FAILED: {error_msg}")
                    await doc_repo.update_status(
                        doc,
                        Document.STATUS_FAILED,
                        error_message=error_msg
                    )
                    return False
            
            except Exception as e:
                logger.warning(f"Error polling parse task {task_id}: {e}")
        
        await doc_repo.update_status(doc, Document.STATUS_FAILED, error_message="Processing timeout")
        return False
    
    async def list_documents(
        self,
        kb_id: str,
        user_id: str,
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[dict], int]:
        """List documents in knowledge base (admin users can access any KB)."""
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)
        
        documents, total = await self.doc_repo.list_documents(kb_id, page, page_size)
        return [doc.to_dict() for doc in documents], total
    
    async def get_document_status(self, doc_id: str, kb_id: str, user_id: str) -> dict:
        """Get document processing status (admin users can access any KB)."""
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)
        
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        return {
            "status": doc.status,
            "errorMessage": doc.error_message,
            "chunkCount": doc.chunk_count
        }
    
    async def get_document_url(self, doc_id: str, kb_id: str, user_id: str) -> dict:
        """Get presigned URL for document file (admin users can access any KB)."""
        from utils.minio_client import get_file_url
        
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)
        
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        if not doc.file_path:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document file not found"}}
            )
        
        # Extract object name from file_path
        object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")
        
        # Generate presigned URL (valid for 1 hour)
        file_url = get_file_url(object_name, expires_seconds=3600)
        
        return {
            "url": file_url,
            "name": doc.name
        }
    
    async def get_document_markdown(self, doc_id: str, kb_id: str, user_id: str) -> str:
        """Get markdown content of a document (admin users can access any KB)."""
        from utils.minio_client import download_file
        
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)
        
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        if not doc.markdown_path:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Markdown content not available"}}
            )
        
        # Extract object name from markdown_path
        object_name = doc.markdown_path.replace(f"{settings.MINIO_BUCKET}/", "")
        
        try:
            # Download markdown from MinIO
            markdown_bytes = await download_file(object_name)
            markdown_content = markdown_bytes.decode('utf-8')
            return markdown_content
        except Exception as e:
            logger.error(f"Failed to get markdown content: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": {"code": "INTERNAL_ERROR", "message": f"Failed to retrieve markdown content: {e}"}}
            )
    
    async def get_documents_markdown_batch(
        self,
        doc_ids: List[str],
        kb_id: str,
        user_id: str
    ) -> dict:
        """
        Batch get markdown content of multiple documents (for agent use).
        
        Args:
            doc_ids: List of document IDs
            kb_id: Knowledge base ID
            user_id: User ID
            
        Returns:
            {
                "documents": {doc_id: markdown_content},
                "document_names": {doc_id: doc_name},  # 🔑 新增：文档名称映射
                "failed": [doc_id]  # IDs that failed to load
            }
        """
        from utils.minio_client import download_file
        
        # Verify access permission (admin users can access any KB)
        await self._verify_kb_access(kb_id, user_id)
        
        logger.info(f"Batch loading markdown for {len(doc_ids)} documents in KB {kb_id}")
        
        documents = {}
        document_names = {}  # 🔑 新增：文档名称映射
        failed = []
        
        # 并发加载所有文档
        async def load_single_doc(doc_id: str):
            try:
                doc = await self.doc_repo.get_by_id(doc_id, kb_id)
                if not doc:
                    logger.warning(f"Document {doc_id} not found")
                    return doc_id, None, None, "not_found"
                
                doc_name = doc.name  # 🔑 获取文档原始名称
                
                if not doc.markdown_path:
                    logger.warning(f"Document {doc_id} has no markdown")
                    return doc_id, None, doc_name, "no_markdown"
                
                # Extract object name from markdown_path
                object_name = doc.markdown_path.replace(f"{settings.MINIO_BUCKET}/", "")
                
                # Download markdown from MinIO
                markdown_bytes = await download_file(object_name)
                markdown_content = markdown_bytes.decode('utf-8')
                
                logger.info(f"Loaded markdown for doc {doc_id} ({len(markdown_content)} chars)")
                return doc_id, markdown_content, doc_name, None
                
            except Exception as e:
                logger.error(f"Failed to load markdown for doc {doc_id}: {e}")
                return doc_id, None, None, str(e)
        
        # 并发执行所有加载任务
        tasks = [load_single_doc(doc_id) for doc_id in doc_ids]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 收集结果
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Task failed with exception: {result}")
                continue
            
            doc_id, content, doc_name, error = result
            
            # 🔑 即使内容加载失败，也保存文档名称（如果有）
            if doc_name:
                document_names[doc_id] = doc_name
            
            if content:
                documents[doc_id] = content
            else:
                failed.append(doc_id)
        
        logger.info(f"Batch load complete: {len(documents)} succeeded, {len(failed)} failed, {len(document_names)} names collected")
        
        return {
            "documents": documents,
            "document_names": document_names,  # 🔑 返回文档名称映射
            "failed": failed
        }
    
    async def retry_document(
        self,
        doc_id: str,
        kb_id: str,
        user_id: str,
        background_tasks
    ) -> dict:
        """
        Retry processing a failed document.
        Only owner and admin users can retry documents.
        
        Smart retry logic:
        - If markdown already exists (e.g., MinerU conversion succeeded but chunking failed),
          skip the conversion step and only retry chunking/embedding
        - Otherwise, re-run the full pipeline
        
        Note: Before retry, we clean up any existing ES data to avoid duplicates.
        """
        # Verify KB write access (owner or admin only)
        kb = await self._verify_kb_write_access(kb_id, user_id)
        
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        # Only allow retry for failed documents
        if doc.status != Document.STATUS_FAILED:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": {"code": "INVALID_STATUS", "message": f"Can only retry failed documents. Current status: {doc.status}"}}
            )
        
        # Check if original file exists
        if not doc.file_path:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": {"code": "FILE_NOT_FOUND", "message": "Original file not found, please re-upload"}}
            )
        
        # Get user's ES index name
        user_es_index = get_user_es_index(user_id)
        
        # Reset document status immediately so UI updates
        await self.doc_repo.update_status(doc, Document.STATUS_PROCESSING, error_message=None)
        
        logger.info(f"Retrying document {doc_id} ({doc.name}) with ES index: {user_es_index}")
        
        # Start background task - all heavy I/O operations happen there
        background_tasks.add_task(
            self._retry_document_background,
            str(doc.id),
            str(doc.kb_id),
            user_es_index,
            doc.file_path,
            doc.markdown_path,
            doc.name
        )
        
        return {
            "id": str(doc.id),
            "name": doc.name,
            "status": Document.STATUS_PROCESSING
        }
    
    async def _retry_document_background(
        self,
        doc_id: str,
        kb_id: str,
        es_index_name: str,
        file_path: str,
        markdown_path: Optional[str],
        filename: str
    ):
        """
        Background task for retry - handles all heavy I/O operations.
        
        This method:
        1. Cleans up existing ES data
        2. Checks if markdown exists (skip conversion if yes)
        3. Downloads necessary files
        4. Calls appropriate processing pipeline
        """
        from config.database import AsyncSessionLocal
        from repositories.document_repository import DocumentRepository
        from utils.minio_client import download_file
        
        async with AsyncSessionLocal() as db:
            doc_repo = DocumentRepository(db)
            
            result = await db.execute(
                select(Document).where(Document.id == doc_id)
            )
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error(f"[Retry] Document {doc_id} not found in background task")
                return
            
            try:
                # Step 1: Clean up existing ES data
                try:
                    await DocumentProcessService.delete_document_from_es(doc_id, es_index_name)
                    logger.info(f"[Retry] Cleaned up existing ES data for document {doc_id}")
                except Exception as e:
                    logger.warning(f"[Retry] Failed to clean up ES data for {doc_id}: {e}")
                
                # Step 2: Check if markdown exists
                has_markdown = False
                markdown_content = None
                if markdown_path:
                    try:
                        md_object_name = markdown_path.replace(f"{settings.MINIO_BUCKET}/", "")
                        markdown_bytes = await download_file(md_object_name)
                        markdown_content = markdown_bytes.decode('utf-8')
                        has_markdown = bool(markdown_content)
                        logger.info(f"[Retry] Document {doc_id} has existing markdown ({len(markdown_content)} chars)")
                    except Exception as e:
                        logger.warning(f"[Retry] Failed to load markdown for {doc_id}: {e}, will re-convert")
                
                # Step 3: Process based on what we have
                if has_markdown:
                    # Skip conversion, only retry chunking
                    await self._retry_chunking_only(doc_id, es_index_name, markdown_content, filename)
                else:
                    # Download original file and run full pipeline
                    try:
                        object_name = file_path.replace(f"{settings.MINIO_BUCKET}/", "")
                        file_data = await download_file(object_name)
                        logger.info(f"[Retry] Downloaded original file for {doc_id}")
                        await self._process_document_pipeline(doc_id, es_index_name, file_data, filename)
                    except Exception as e:
                        logger.error(f"[Retry] Failed to download original file for {doc_id}: {e}")
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=f"Failed to download original file: {e}"
                        )
                        
            except Exception as e:
                logger.error(f"[Retry] Background task failed for {doc_id}: {e}")
                await doc_repo.update_status(
                    doc,
                    Document.STATUS_FAILED,
                    error_message=str(e)
                )
    
    async def _retry_chunking_only(
        self,
        doc_id: str,
        es_index_name: str,
        markdown_content: str,
        filename: str
    ):
        """
        Retry only the chunking/embedding step when markdown already exists.
        This is much faster than re-running MinerU conversion.
        """
        from config.database import AsyncSessionLocal
        from repositories.document_repository import DocumentRepository
        from repositories.kb_repository import KnowledgeBaseRepository
        
        async with AsyncSessionLocal() as db:
            doc_repo = DocumentRepository(db)
            kb_repo = KnowledgeBaseRepository(db)
            
            result = await db.execute(
                select(Document).where(Document.id == doc_id)
            )
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error(f"Document {doc_id} not found in retry task")
                return
            
            logger.info(f"[Retry] Starting chunking-only retry for {doc_id} ({filename})")
            
            try:
                # Go directly to chunking step
                await doc_repo.update_status(doc, Document.STATUS_CHUNKING)
                
                # Save markdown to temp file for processing
                temp_file_path = f"/tmp/{doc_id}.md"
                with open(temp_file_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                parse_filename = os.path.splitext(filename)[0] + '.md'
                
                try:
                    logger.info(f"[Retry] Calling document processing service for {doc_id}...")
                    parse_result = await DocumentProcessService.parse_document(
                        temp_file_path,
                        str(doc_id),
                        es_index_name,
                        parse_filename
                    )
                    task_id = parse_result["task_id"]
                    logger.info(f"[Retry] Parse task created: {task_id}")
                    
                    await doc_repo.update_status(
                        doc,
                        Document.STATUS_EMBEDDING,
                        parse_task_id=task_id
                    )
                    
                    # Poll parsing status
                    success = await self._poll_parse_task(doc, task_id, doc_repo)
                    
                    if success:
                        logger.info(f"[Retry] Document {doc_id} processing completed successfully!")
                        await kb_repo.increment_contents_count(doc.kb_id)
                    else:
                        logger.warning(f"[Retry] Document {doc_id} processing failed")
                    
                finally:
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                        
            except Exception as e:
                logger.error(f"[Retry] Chunking retry failed for {doc_id}: {e}")
                await doc_repo.update_status(
                    doc,
                    Document.STATUS_FAILED,
                    error_message=str(e)
                )
    
    async def delete_document(self, doc_id: str, kb_id: str, user_id: str):
        """
        Delete document from KB, MinIO, and ES.
        Only owner and admin users can delete documents.
        """
        # Verify KB write access (owner or admin only)
        kb = await self._verify_kb_write_access(kb_id, user_id)
        
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        # Get user's ES index name
        user_es_index = get_user_es_index(user_id)
        
        # Delete from ES (using user-level index)
        if doc.status == Document.STATUS_READY:
            try:
                await DocumentProcessService.delete_document_from_es(str(doc.id), user_es_index)
            except Exception as e:
                logger.warning(f"Failed to delete from ES: {e}")
        
        # Delete from MinIO (original file)
        if doc.file_path:
            try:
                object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")
                await delete_file(object_name)
            except Exception as e:
                logger.warning(f"Failed to delete from MinIO: {e}")
        
        # Delete markdown from MinIO
        if doc.markdown_path:
            try:
                md_object_name = doc.markdown_path.replace(f"{settings.MINIO_BUCKET}/", "")
                await delete_file(md_object_name)
            except Exception as e:
                logger.warning(f"Failed to delete markdown from MinIO: {e}")
        
        # Delete from DB
        await self.doc_repo.delete(doc)
        
        # Decrement KB contents count (only for successfully processed documents)
        # 只有成功处理的文档才会在处理完成时增加计数，所以删除时也只减少成功的文档
        if doc.status == Document.STATUS_READY:
            await self.kb_repo.increment_contents_count(kb_id, -1)
            logger.info(f"Deleted document: {doc_id} (decremented contents count)")
        else:
            logger.info(f"Deleted document: {doc_id} (status: {doc.status}, no count change)")
    
    async def move_document(
        self,
        doc_id: str,
        source_kb_id: str,
        target_kb_id: str,
        user_id: str
    ) -> dict:
        """
        Move document from one knowledge base to another.
        Only owner can move documents between their own knowledge bases.
        
        Note: Since ES index is user-level (not KB-level), no ES migration needed.
        MinIO file paths also don't need to change as they're stored by user_id.
        """
        # Verify write access to both KBs (must be owner of both)
        source_kb = await self._verify_kb_write_access(source_kb_id, user_id)
        target_kb = await self._verify_kb_write_access(target_kb_id, user_id)
        
        # Get document
        doc = await self.doc_repo.get_by_id(doc_id, source_kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": {"code": "NOT_FOUND", "message": "Document not found"}}
            )
        
        # Cannot move to the same KB
        if source_kb_id == target_kb_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": {"code": "INVALID_REQUEST", "message": "Source and target knowledge base are the same"}}
            )
        
        # Move document (update kb_id)
        await self.doc_repo.update_kb_id(doc, target_kb_id)
        
        # Update contents count for both KBs (only for ready documents)
        if doc.status == Document.STATUS_READY:
            await self.kb_repo.increment_contents_count(source_kb_id, -1)
            await self.kb_repo.increment_contents_count(target_kb_id, 1)
            logger.info(f"Moved document {doc_id} from KB {source_kb_id} to KB {target_kb_id}")
        else:
            logger.info(f"Moved document {doc_id} (status: {doc.status}, no count change)")
        
        return {
            "id": str(doc.id),
            "name": doc.name,
            "sourceKbId": source_kb_id,
            "targetKbId": target_kb_id,
            "status": doc.status
        }

